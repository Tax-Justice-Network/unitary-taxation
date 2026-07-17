"""
9r — standalone TABLES-chapter document.

Builds a separate .docx containing the paper's Tables 1-4 in the same rendering
style as the autodraft (grey bold header, right-aligned numbers, shaded group
rows — reuses _build_paper_docx helpers), each with a self-explanatory caption
note in the house style of the figure notes (user 2026-07-12).

Output: "<paper folder>/Unitary taxation_tables_vN.docx" (next free N).
Usage:  python 9r_tables_document.py
"""
import os
import sys

import docx
from docx.shared import Pt

_SRC = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _SRC)

from _build_paper_docx import (  # noqa: E402
    TBL, _DIR, add_par_after, add_table_from_csv_after,
)

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# ── shared caption fragments (house style, mirroring the figure notes) ───────
_BASIS = (" The estimates are based on aggregate country by country reporting data "
          "from the OECD. Because the estimates are based only on the sample of "
          "multinationals for which data of sufficient quality are available in the "
          "OECD data, the values represent only a part of the total expected gains "
          "or losses. Values are annual averages over 2016–2022, excluding 2020, in "
          "billions of constant 2025 US dollars. Group rows (in capitals) follow "
          "the World Bank income classification, with tax havens shown "
          "separately.")
_FORMULAS = (" Four apportionment formulas are shown: sales & employees — our "
             "preferred formula — weighting destination-based sales and employee "
             "headcount 50 per cent each; the formula of the EU's Common "
             "Consolidated Corporate Tax Base (one third sales, one third assets, "
             "one sixth headcount, one sixth payroll); a three-factor formula (one "
             "third each sales, assets and headcount); and a double-weighted sales "
             "formula (50 per cent sales, 25 per cent assets and headcount each). "
             "In all formulas, sales are measured at destination.")
_RATES = (" Taxable profits gained under unitary taxation are assumed to be taxed "
          "at the statutory corporate income tax rate; profits lost are assumed to "
          "have previously been taxed at the effective tax rate applying to the "
          "profits concerned.")

TABLES = [
    ("Table 1: Change in taxable profits",
     "Notes: This table reports the estimated change in each country's taxable "
     "profits under unitary taxation, with resource-related profits excluded from "
     "the unitary taxation profit base (resource rights prior to taxing rights). "
     "The first data column gives each country's current profit base (the sum of "
     "positive reported profits); columns marked '(%)' express the change relative "
     "to it and are not meaningful where that base is zero."
     + _FORMULAS + _BASIS,
     [(None, "table2a_taxbase_by_formula__reported_only.csv")]),
    ("Table 2: Revenue estimates from switching to unitary taxation",
     "Notes: This table reports the estimated change in each country's corporate "
     "income tax revenues under unitary taxation, with resource-related profits "
     "excluded from the unitary taxation profit base. The first data column gives "
     "the corporate tax each country currently collects from the multinationals "
     "in the data; columns marked '(%)' express the change relative to it."
     + _RATES + _FORMULAS + _BASIS,
     [(None, "table2b_revenue_by_formula__reported_only.csv")]),
    # (block label below matches the docx builder's Table 3 sub-labels)
    ("Table 3: Break-even ETRs of losers from unitary taxation",
     "Notes: This table reports the effective tax rates at which countries losing "
     "tax revenue under unitary taxation would preserve their current tax revenues "
     "from multinationals. In both blocks, the break-even rate is the rate the "
     "country would need to apply to the entire profit base that unitary taxation "
     "allocates to it under each formula — including, for the home-bias countries "
     "in the second block, the enlarged bases of foreign multinationals. The "
     "descriptive columns report the rates multinationals currently face: for tax "
     "havens the average effective rate, and for home-bias countries the domestic "
     "effective rate (which explains why they lose), the effective rate on all "
     "multinational profit, and the statutory rate. Rates in per cent, pooled over "
     "2016–2022 excluding 2020; preferred destination-based sales measure "
     "throughout. The figures for Japan and Saudi Arabia are affected by the "
     "imperfect removal of resource-related taxation — see the main text." + _BASIS.replace(
         " Values are annual averages over 2016–2022, excluding 2020, in billions "
         "of constant 2025 US dollars.", ""),
     [("Tax havens (average ETR):",
       "table5_breakeven_etr__reported_only.csv"),
      ("Home-bias countries (domestic ETR):",
       "table6_breakeven_domestic_etr__reported_only.csv")]),
    ("Table 4: Unitary taxation effects with and without resource rent correction",
     "Notes: This table reports each country's reported profit base and current "
     "tax revenue from multinationals in the data, followed by the estimated "
     "changes in taxable profits and tax revenues under three treatments of the "
     "resource sector: ignoring resource rights (baseline), granting resource "
     "rights priority over taxing rights (resource-related profits excluded — the "
     "main specification), and additionally enforcing the minimum royalty (total "
     "revenue change, with the royalty component shown separately in the 'Of "
     "which: royalty' column). The final column adjusts the main specification "
     "for cross-border loss consolidation (a lower-bound revenue estimate). "
     "Preferred sales & employees formula with destination-based sales."
     + _RATES + _BASIS,
     [(None, "table1_scenarios__reported_only.csv")]),
]


def main():
    target = None
    for n in range(1, 40):
        cand = os.path.join(_DIR, f"Unitary taxation_tables_v{n}.docx")
        if not os.path.exists(cand):
            target = cand
            break
    doc = docx.Document()
    h = doc.add_heading("TABLES", level=1)
    cur = h._p
    for title, note, items in TABLES:
        hd = doc.add_heading(title, level=2)
        cur.addnext(hd._p)
        cur = hd._p
        np_ = doc.add_paragraph()
        r = np_.add_run(note)
        r.italic = True
        r.font.size = Pt(8)
        cur.addnext(np_._p)
        cur = np_._p
        for label, fname in items:
            if label:
                cur = add_par_after(doc, cur, label, bold=True)
            csv_path = os.path.join(TBL, fname)
            if not os.path.exists(csv_path):
                cur = add_par_after(doc, cur, f"[missing: {fname}]")
                continue
            el, nrow = add_table_from_csv_after(doc, cur, csv_path)
            print(f"  + {title[:34]} <- {fname} ({nrow} rows)")
            cur = el
    doc.save(target)
    print(f"\nSaved: {target}")


if __name__ == "__main__":
    main()
