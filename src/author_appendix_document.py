"""
9s — standalone document with Appendix F, G and H.

  Appendix F: Results with disaggregated rows (gravity sample) — intro prose,
              Tables F1–F3, Figures F1–F7 with house-style notes. Filled only
              when the gravity passes (stage 2d) have completed; otherwise the
              slots carry [pending] markers and the script is simply rerun later.
  Appendix G: Tax revenue estimates using ETR for gains and losses — Figure G1
              (the ETR–ETR robustness figure) with note.
  Appendix H: Overview of hand-collected data sources — prose + table.

Output: "<paper folder>/Unitary taxation_appendixFGH_vN.docx" (next free N).
Usage:  python 9s_appendix_document.py
"""
import os
import sys

import docx
from docx.shared import Pt

_SRC = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(_SRC)
sys.path.insert(0, _SRC)

from _build_paper_docx import (  # noqa: E402
    TBL, _DIR, FIG_GRAV, FIG_REP, add_par_after, add_table_from_csv_after,
    TXT_APP_E, TXT_APP_F,
)

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

_CMN = (" The estimates are based on aggregate country by country reporting data "
        "from the OECD. Income groups follow the World Bank classification, with "
        "tax havens shown separately. Values are annual averages over 2016–2022, "
        "excluding 2020, in billions of constant 2025 US dollars.")
_GRAV = ("Notes: Disaggregated (gravity-imputed) sample — the aggregated "
         "rest-of-world and continental rows of limited-detail reporters are "
         "split across individual partner jurisdictions using the gravity-based "
         "imputation of Appendix E. ")

# Ordered to mirror the MAIN report's table sequence (user 2026-07-13):
# Table 1 = taxbase by formula, Table 2 = revenue by formula, Table 4 =
# scenario comparison (there is no gravity Table-3 / break-even analogue).
APPENDIX_F_TABLES = [
    ("Table F1. Change in taxable profit by apportionment formula, resources "
     "excluded (disaggregated / gravity sample)",
     "table2a_taxbase_by_formula__gravity.csv"),
    ("Table F2. Change in tax revenue by apportionment formula, resources "
     "excluded (disaggregated / gravity sample)",
     "table2b_revenue_by_formula__gravity.csv"),
    ("Table F3. Change in taxable profit and tax revenue by scenario "
     "(disaggregated / gravity sample; sales & employees)",
     "table1_scenarios__gravity.csv"),
]
# (label, gravity png, one-line description for the note — number-free so the
# author's main-text renumbering cannot go stale)
APPENDIX_F_FIGURES = [
    ("Figure F1. Change in taxable profits under unitary taxation, by income group "
     "(resources excluded)", "fig01_taxable_profit_by_income.png",
     "Mirrors the main text's change-in-taxable-profits figure."),
    ("Figure F2. Change in tax revenue by income group (resources excluded; gains "
     "at CIT, losses at matched ETR)", "fig02_tax_revenue_ETR_CIT.png",
     "Mirrors the main text's change-in-tax-revenue figure."),
    ("Figure F3. Change in taxable profit across the three resource treatments "
     "(sales & employees)", "fig04_taxbase_three_treatments.png",
     "Taxable-profit companion of the main text's resource-treatment revenue figure."),
    ("Figure F4. Change in tax revenue across the three resource treatments "
     "(sales & employees)", "fig06_revenue_three_treatments.png",
     "Mirrors the main text's resource-treatment revenue figure."),
    ("Figure F5. Change in taxable profits by income group: origin- versus "
     "destination-based sales", "fig08_origin_dest_nexus_taxbase.png",
     "Taxable-profit companion of the main text's origin-versus-destination revenue figure."),
    ("Figure F6. Change in tax revenue by income group: origin- versus "
     "destination-based sales", "fig09_origin_dest_revenue.png",
     "Mirrors the main text's origin-versus-destination revenue figure."),
    ("Figure F7. Change in taxable profits under alternative destination-based "
     "sales measures", "figA_sales_measure_ablation.png",
     "Compares the complete destination measure, the variant without digital "
     "services, and the consumer-facing-only variant."),
]

TXT_APP_G = [
    "Our headline revenue estimates value the profits a country gains under "
    "unitary taxation at its statutory corporate income tax rate, and the profits "
    "it loses at the effective tax rate applying to the profits concerned. This "
    "appendix reports the more conservative specification in which BOTH gains and "
    "losses are valued at effective tax rates: newly allocated profits are assumed "
    "to be taxed at the rate the country currently applies to comparable profits, "
    "rather than at the statutory rate. Every revenue figure of the main text is "
    "reproduced below under this specification; taxable-profit figures are "
    "unaffected by the rate choice.",
]
_ETRNOTE = ("Notes: As the corresponding main-text figure, but with both the "
            "profits a country gains and the profits it loses under unitary "
            "taxation valued at effective tax rates (conservative rate "
            "specification). Resource-related profits are excluded from the "
            "unitary taxation profit base unless the figure compares resource "
            "treatments.")
# (label, file, extra note sentence)
APPENDIX_ETR_FIGURES = [
    ("Figure G1. Global revenue gains from unitary taxation (gains and losses at "
     "effective tax rates)", "fig_scaleup_gross_yearly_etretr.png",
     " Gross gains of winning jurisdictions; scaled bars as in the main text."),
    ("Figure G2. Global revenue gains net of losses (gains and losses at "
     "effective tax rates)", "fig_scaleup_net_yearly_etretr.png",
     " Net of the losses of tax havens and other losing jurisdictions."),
    ("Figure G3. Change in tax revenue by income group (resources excluded; gains "
     "and losses at effective tax rates)", "fig03_tax_revenue_ETR_ETR.png",
     " The percentage panel expresses the change relative to each group's current "
     "tax revenues from multinationals in the data."),
    ("Figure G4. Change in tax revenue across the three resource treatments "
     "(gains and losses at effective tax rates)",
     "figE_revenue_three_treatments_etretr.png",
     " The hatched slice is the minimum-royalty floor add-on."),
    ("Figure G5. Angola and Peru (gains and losses at effective tax rates)",
     "figE_country_angola_peru_etretr.png", ""),
    ("Figure G6. Guinea and Mali (gains and losses at effective tax rates)",
     "figE_country_gin_mli_etretr.png", ""),
    ("Figure G7. Change in tax revenue by income group: origin- versus "
     "destination-based sales (gains and losses at effective tax rates)",
     "figE_origin_dest_revenue_etretr.png", ""),
]


def _gravity_ready():
    # true readiness = the gravity FLOORED scenario has completed (its summary
    # exists) AND the gravity tables were rebuilt afterwards.
    floored_summary = os.path.join(
        _ROOT, "output", "unitary_taxation", "gravity", "excl_resource_floored",
        "tables", "summary_country_year_long.csv")
    return (os.path.exists(floored_summary)
            and all(os.path.exists(os.path.join(TBL, f))
                    for _, f in APPENDIX_F_TABLES))


def add_note_after(doc, cur, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    cur.addnext(p._p)
    return p._p


def add_image_after_el(doc, cur, img_path, width_in=6.5):
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    cur.addnext(p._p)
    return p._p


def main():
    target = None
    for n in range(1, 40):
        cand = os.path.join(_DIR, f"Unitary taxation_appendixFGH_v{n}.docx")
        if not os.path.exists(cand):
            target = cand
            break
    doc = docx.Document()

    # ── Appendix F ───────────────────────────────────────────────────────────
    h = doc.add_heading("Appendix F: Results with disaggregated rows", level=1)
    cur = h._p
    for t in TXT_APP_E:
        cur = add_par_after(doc, cur, t)
    ready = _gravity_ready()
    if not ready:
        cur = add_par_after(
            doc, cur, "[PENDING — the full-sample (gravity) estimation has not "
            "completed yet; rerun 9s_appendix_document.py after stage 2d and the "
            "gravity 9j/9m passes to fill the tables and figures below.]", bold=True)
    # Figures BEFORE tables (user 2026-07-13).
    for title, fname, desc in APPENDIX_F_FIGURES:
        hd = doc.add_heading(title, level=2)
        cur.addnext(hd._p)
        cur = hd._p
        fp = os.path.join(FIG_GRAV, fname)
        if ready and os.path.exists(fp):
            cur = add_image_after_el(doc, cur, fp)
            cur = add_note_after(doc, cur, _GRAV + desc + _CMN)
            print(f"  + {fname}")
        else:
            cur = add_par_after(doc, cur, f"[pending: {fname}]")
    for title, fname in APPENDIX_F_TABLES:
        hd = doc.add_heading(title, level=2)
        cur.addnext(hd._p)
        cur = hd._p
        fp = os.path.join(TBL, fname)
        if ready and os.path.exists(fp):
            el, nrow = add_table_from_csv_after(doc, cur, fp)
            print(f"  + {fname} ({nrow} rows)")
            cur = el
        else:
            cur = add_par_after(doc, cur, f"[pending: {fname}]")

    # ── ETR-ETR appendix (Appendix G in the author's lettering, 2026-07-13) ──
    h = doc.add_heading("Appendix G: Tax revenue estimates using ETR for gains "
                        "and losses", level=1)
    cur = h._p
    for t in TXT_APP_G:
        cur = add_par_after(doc, cur, t)
    for g_title, g_file, extra in APPENDIX_ETR_FIGURES:
        hd = doc.add_heading(g_title, level=2)
        cur.addnext(hd._p)
        cur = hd._p
        gp = os.path.join(FIG_REP, g_file)
        if os.path.exists(gp):
            cur = add_image_after_el(doc, cur, gp)
            cur = add_note_after(doc, cur, _ETRNOTE + extra + _CMN)
            print(f"  + {g_file}")
        else:
            cur = add_par_after(doc, cur, f"[missing: {g_file}]")

    # ── Appendix H ───────────────────────────────────────────────────────────
    h = doc.add_heading("Appendix H: Overview of hand-collected data sources", level=1)
    cur = h._p
    for t in TXT_APP_F:
        cur = add_par_after(doc, cur, t)
    hp = os.path.join(TBL, "appendixF_handcollected_sources.csv")
    if os.path.exists(hp):
        hd = doc.add_heading("Table H1. Overview of hand-collected data sources",
                             level=2)
        cur.addnext(hd._p)
        cur = hd._p
        el, nrow = add_table_from_csv_after(doc, cur, hp)
        print(f"  + hand-collected sources table ({nrow} rows)")
    else:
        cur = add_par_after(doc, cur, "[missing: appendixF_handcollected_sources.csv]")

    doc.save(target)
    print(f"\nSaved: {target}")


if __name__ == "__main__":
    main()
