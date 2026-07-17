"""
_build_paper_docx.py — assemble a fresh COPY of the unitary-taxation paper with the
result/appendix FIGURES + country TABLES embedded, the empty Appendix C/E/F bodies
written, and the reference list expanded. The MAIN document is only read, never
modified. Output goes to the next free `Unitary taxation_autodraft_vN.docx`.

Inserts (sourced from the clean main doc each run, so a rebuild is a fresh copy):
  • Main-text figures: the author now owns caption + note TEXT in the main doc
    (draft renumbered 2026-07-12) — the builder only REFRESHES the image at each
    caption, anchored by caption text (FIG_CAPTIONS), never by figure number.
  • Tables 1–4 (reported) under their headings in the "Tables" section.
  • Appendix F (gravity results): tables under the Table F1–F3 headings, images
    under the Figure F1–F7 headings (F7 = sales-measure ablation, retitled).
  • Appendix D dividend prose / Appendix F intro / Appendix G prose+table are
    inserted ONLY if not already merged into the main doc (guarded by phrase).
  • Expanded references (flagged block) after the References heading.

Does NOT touch the methodology text (already pasted into the main doc).
Usage: python _build_paper_docx.py
"""
import os
import sys
import shutil
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

_SRC = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(_SRC)
sys.path.insert(0, _SRC)
import config  # noqa: E402

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

_DIR = r"C:\Users\aliso\Tax Justice Network Ltd\TJN - Shared Documents\Research team\Projects long-term\Unitary taxation"
ORIG = os.path.join(_DIR, "Unitary taxation.docx")          # the MAIN doc (read-only source)

MARK = "⟦INSERTED BY SCRIPT — review / renumber / delete this marker⟧"

TBL = os.path.join(_ROOT, "output", "unitary_taxation", "across_samples",
                   "paper_tables", "tables")
FIG_REP = os.path.join(_ROOT, "output", "unitary_taxation", "reported_only",
                       "paper_figures", "figures")
FIG_GRAV = os.path.join(_ROOT, "output", "unitary_taxation", "gravity",
                        "paper_figures", "figures")

# Main-text figures, anchored by a NUMBER-FREE caption fragment (the author
# renumbers freely; find_caption requires the paragraph to start with "Figure"
# and contain the fragment). The builder replaces only the image sitting
# at/above each caption — the caption and any note text stay exactly as the
# author wrote them.
FIG_CAPTIONS = [
    ("Global revenue gains from unitary taxation",          # Fig 1 precedes Fig 2 in doc order
     "fig_scaleup_gross_yearly.png"),
    ("net of losses",
     "fig_scaleup_net_yearly.png"),
    ("Change in taxable profits under unitary taxation",
     "fig01_taxable_profit_by_income.png"),
    ("Change in corporate income tax revenue under unitary taxation",
     "fig02_tax_revenue_ETR_CIT.png"),
    ("Estimated change in tax revenue with different resource treatment",
     "fig06_revenue_three_treatments.png"),
    ("Angola and Peru",
     "fig05_country_angola_peru.png"),
    ("Guinea and Mali",
     "fig07_country_gin_mli.png"),
    ("Change in taxable profits by income group: origin",
     "fig08_origin_dest_nexus_taxbase.png"),
    ("Change in tax revenue by income group: origin",
     "fig09_origin_dest_revenue.png"),
    ("alternative destination-based sales measures",
     "figA_sales_measure_ablation.png"),
    ("cross-border loss consolidation",
     "fig11_loss_consolidation.png"),
    # ETR-ETR robustness (labelled "Figure 3." again in its own section).
    ("Change in corporate income tax revenue by income group",
     "fig03_tax_revenue_ETR_ETR.png"),
    # Factor incidence (§4 draft: "How each factor benefits or harms different
    # income groups") — warns until the author adds the caption to the doc.
    ("How each factor benefits or harms",
     "fig10_factor_incidence.png"),
]
# Appendix F (results on the disaggregated / gravity sample): images live under
# their own Heading-3 slots. F7 is the sales-measure ablation; its stale heading
# text (old CFB variants) is retitled on insert.
FIG_APPENDIX_F = [
    ("Figure F1.", "fig01_taxable_profit_by_income.png", None),
    ("Figure F2.", "fig02_tax_revenue_ETR_CIT.png", None),
    ("Figure F3.", "fig04_taxbase_three_treatments.png", None),
    ("Figure F4.", "fig06_revenue_three_treatments.png", None),
    ("Figure F5.", "fig08_origin_dest_nexus_taxbase.png", None),
    ("Figure F6.", "fig09_origin_dest_revenue.png", None),
    ("Figure F7.", "figA_sales_measure_ablation.png",
     "Figure F7. Change in taxable profits under alternative destination-based "
     "sales measures"),
]

TXT_APP_C = [
    "The OECD country-by-country reporting data are known to double-count intra-group "
    "dividends: profit earned in an affiliate and then distributed up the ownership chain is "
    "recorded a second time as profit in the receiving (often haven) jurisdiction. This "
    "inflates both the profit booked in tax havens and the global profit total. We remove "
    "this double counting following García-Bernardo, Janský and Zucman (2026).",
    "Procedure. For each partner jurisdiction j classified as a tax haven (list below), and "
    "for non-US-headquartered multinationals only, we reduce reported profit before tax by "
    "10% in each year 2016–2019; the associated tax is left unchanged. That is, "
    "profit_cleaned_ijt = profit_reported_ijt × (1 − 0.10) when j is a haven, the parent i is "
    "non-US, and t ∈ {2016, 2017, 2018, 2019}, and profit_cleaned = profit_reported otherwise. "
    "The 10% factor is the share that García-Bernardo, Janský and Zucman (2026) estimate to be "
    "intra-group dividends over that period; restricting the correction to non-US parents and "
    "to 2016–2019 follows their scope (US multinationals report on a basis less affected by "
    "this artefact, and the estimate is identified over 2016–2019). The adjustment lowers "
    "reported haven profit and the global total; all downstream estimates use the cleaned series.",
    "Tax-haven list. The paper does not otherwise take a stance on which jurisdictions are "
    "\"tax havens\"; for this correction we use exactly the list of García-Bernardo, Janský and "
    "Zucman (2026), grouped following Reurink and García-Bernardo (2020): eleven \"profit "
    "centres\" — the Bahamas, Bermuda, Barbados, the Cayman Islands, Gibraltar, the Isle of Man, "
    "Jersey, Malta, Mauritius, Puerto Rico and the British Virgin Islands — and six "
    "\"coordination centres\" — Switzerland, Hong Kong, Ireland, Luxembourg, the Netherlands and "
    "Singapore (17 jurisdictions in total).",
    "Uncertainties. The correction is deliberately simple and carries several caveats. (i) The "
    "10% rate is a single pooled estimate applied uniformly across havens and years, whereas the "
    "true dividend share varies by jurisdiction, parent and year. (ii) The haven list is a "
    "modelling choice — a broader or narrower list removes more or less profit (we use the "
    "García-Bernardo–Janský–Zucman list here for comparability, and a separate, wider list only "
    "for the presentational income-group classification). (iii) The scope restriction (non-US "
    "parents, 2016–2019) leaves later years and US-parented groups uncorrected, so residual "
    "double counting may remain there. (iv) The adjustment scales profit proportionally and does "
    "not trace individual dividend flows, which the aggregated data do not permit. These choices "
    "follow the source study; a different rate, list or scope can be substituted and the cleaning "
    "re-run.",
]
TXT_APP_E = [
    "This appendix reproduces the main-text results on the disaggregated (gravity-imputed) "
    "sample, in which the aggregated 'rest-of-world' and continental rows of limited-detail "
    "reporters are split across individual partner jurisdictions using the gravity/"
    "machine-learning imputation described in Appendix D, rather than restricted to directly "
    "reported rows.",
    "The results should be read alongside the reported-sample results in Section 5. The "
    "imputed sample achieves fuller country coverage but carries additional uncertainty on "
    "the imputed cells (quantified by the bootstrap standard errors); for a small number of "
    "micro-states the imputation inflates activity implausibly, and for those jurisdictions "
    "the reported-sample figures are preferred. The tables and figures below mirror Tables "
    "1–2b and the corresponding results figures in the main text.",
]
TXT_APP_F = [
    "This appendix lists the resource-revenue figures, resource profit-tax rates, and "
    "small-territory macro, corporate-income-tax and wage values that were collected by "
    "hand — from EITI reports, company filings and government sources, and, where noted, "
    "with AI-assisted extraction — rather than taken from a single machine-readable dataset.",
    "Each row gives the jurisdiction, the commodity or variable, a confidence rating, and "
    "the source. These values feed the resource-rent-capture correction (Section 4.1) and "
    "the small-territory imputations (Section 3.5).",
]

REFS_ADD = [
    "— Works cited in the text but missing from the list above (please verify details):",
    "Intergovernmental Forum on Mining, Minerals, Metals and Sustainable Development (IGF) "
    "& African Tax Administration Forum (ATAF) (2022). The Future of Resource Taxation: "
    "Variable royalty regimes. IGF/ATAF. [verify exact title & publisher]",
    "Loretz, S. (2026). Revenue effects of unitary taxation with formulary apportionment "
    "(firm-level estimates). WIFO Working Paper. [verify title, number & year]",
    "Palanský, M., & Schultz, A. (2024). [EU-wide revenue effects of formulary "
    "apportionment]. [verify full citation]",
    "OECD (2020). Tax Challenges Arising from the Digitalisation of the Economy — Economic "
    "Impact Assessment. OECD/G20 Inclusive Framework on BEPS, OECD Publishing, Paris. "
    "[the market-allocation method in §4.2.3 is from this report, not the Pillar One "
    "Blueprint — check which OECD (2020) each in-text citation refers to]",
    "Reurink, A., & García-Bernardo, J. (2020). Competing for capital: the great "
    "fragmentation of the firm and the rise of tax haven use. [verify full citation]",
    "— Data sources used:",
    "Bureau van Dijk / Moody's Analytics. Orbis database (company ownership & financials).",
    "Extractive Industries Transparency Initiative (EITI). Summary Data. https://eiti.org",
    "International Labour Organization (ILO). ILOSTAT — wages and working-time statistics.",
    "International Telecommunication Union (ITU). World Telecommunication/ICT Indicators "
    "Database (internet penetration).",
    "OECD. Analytical AMNE database (Activity of Multinational Enterprises).",
    "United Nations Statistics Division. National Accounts Main Aggregates Database "
    "(household final consumption expenditure).",
    "UNU-WIDER (2023). Government Revenue Dataset (GRD).",
    "World Bank. World Development Indicators (GDP, population).",
    "World Health Organization (WHO). Global Health Expenditure Database.",
    "WTO–UNCTAD–ITC. Digitally deliverable services trade dataset (imports of digitally "
    "delivered services).",
]

CAP = {
    "t1_rep": "table [ins.]. change in taxable profit and tax revenue by scenario "
              "(reported sample; sales & employees). US$ bn, 2016–2022; tax revenue on the "
              "ETR-CIT specification (gains at CIT, losses at ETR). The minimum-royalty column "
              "gives total revenue (UT + royalty), with the royalty benefit shown separately "
              "('of which: royalty'); loss consolidation changes only tax revenue.",
    "t2a_rep": "table [ins.]. change in taxable profit by apportionment formula, resources "
               "excluded (reported sample; US$ bn, 2016–2022).",
    "t2b_rep": "table [ins.]. change in tax revenue (ETR-CIT) by apportionment formula, "
               "resources excluded (reported sample; US$ bn, 2016–2022).",
    "t1_grv": "table E1. change in taxable profit and tax revenue by scenario "
              "(disaggregated / gravity sample; sales & employees; US$ bn, 2016–2022; "
              "ETR-CIT).",
    "t2a_grv": "table E2a. change in taxable profit by apportionment formula, resources "
               "excluded (disaggregated / gravity sample; US$ bn, 2016–2022).",
    "t2b_grv": "table E2b. change in tax revenue (ETR-CIT) by apportionment formula, "
               "resources excluded (disaggregated / gravity sample; US$ bn, 2016–2022).",
    "t3a": "table [ins.]. change in taxable profit by sales measure — employees + sales base "
           "(reported sample; US$ bn, 2016–2022; resources excluded).",
    "t3b": "table [ins.]. change in tax revenue (ETR-CIT) by sales measure — employees + sales "
           "base (reported sample; US$ bn, 2016–2022; resources excluded).",
    "t5": "table [ins.]. break-even ETR for investment-hub losers — the effective rate needed to "
          "keep current MNE-tax revenue on the smaller unitary-taxation base, by formula "
          "(reported sample; %; 2016–2022).",
    "t6": "table [ins.]. break-even rates for home-bias losers, by formula (reported sample; %; "
          "2016–2022). Left block: the rate on the domestic cell alone that keeps current "
          "domestic MNE tax revenue — illustrative, as a statutory rate cannot be ring-fenced "
          "to domestic multinationals. Right block: the rate on all profit unitary taxation "
          "allocates to the country (including foreign multinationals' enlarged bases) that "
          "keeps all current MNE tax revenue — the policy-relevant break-even. Japan and Saudi "
          "Arabia are affected by imperfect resource correction — see text.",
    "fF": "table G1. overview of hand-collected data sources.",
    "fig11": "figure 11. revenue effect of cross-border loss consolidation, by income group "
             "(reported sample; sales & employees; resources excluded; ETR-CIT; 2016–2022). "
             "The consolidation-adjusted bar is a lower-bound revenue estimate — it credits "
             "zero tax on reported losses; the consolidation loss itself is an upper bound.",
    "figE": "figure {lab}. disaggregated (gravity) sample — mirrors the corresponding "
            "main-text results figure.",
}


# ── Per-figure title + note REFERENCE TEXT ───────────────────────────────────
# No longer auto-inserted (the author owns caption/note text in the main doc since
# the 2026-07-12 renumbering); kept as the canonical wording to copy from.
_CMN = (" Reported sample, 2016–2022. USD billions; income groups follow the World Bank "
        "classification, with investment hubs shown separately.")
_POSNOTE = (" The percentage panel expresses the change as a share of the group's current-system "
            "profit base, computed over the positive-base subsample (groups whose aggregate "
            "current profit base is non-positive are omitted, as a percentage against a "
            "non-positive base flips sign).")
_REVPCT = (" The percentage panel expresses the change as a share of the corporate tax the "
           "group's reporting multinationals currently pay (cash basis).")
FIG_META = {
    1: ("Change in taxable profits under unitary taxation, by income group (resources excluded)",
        "Four apportionment formulas; sales & employees is the preferred (destination-based) "
        "formula." + _POSNOTE + _CMN),
    2: ("Change in tax revenue by income group (resources excluded; gains at CIT, losses at ETR)",
        "Gains valued at the statutory corporate income tax rate, losses at the effective tax "
        "rate — our preferred specification, used throughout the resource-treatment and "
        "country figures. Four formulas; sales & employees preferred." + _REVPCT + _CMN),
    3: ("Change in tax revenue by income group (resources excluded; gains and losses at ETR)",
        "Robustness to the rate specification: both the profit a country loses and the profit "
        "it gains are valued at the effective tax rate (ETR). Four formulas; sales & employees "
        "preferred." + _REVPCT + _CMN),
    4: ("Change in taxable profit across the three resource treatments (sales & employees)",
        "Resource rights prior to taxing rights (the main specification), the same plus a "
        "minimum royalty, and ignoring resource rights, under the preferred sales & employees "
        "formula. Between-scenario differences are far larger at the country level than the "
        "group totals suggest — e.g. Saudi Arabia's taxable base falls by about US$100 billion "
        "per year when resource rights are ignored versus US$8 billion under the main "
        "specification — because the reallocated resource base largely moves between countries "
        "within the same income group." + _POSNOTE + _CMN),
    5: ("Angola and Peru: change in tax revenue, with and without resource rights "
        "prior to taxing rights",
        "Both are resource-dependent producers (oil; copper and other minerals). When resource "
        "rights are ignored they lose from unitary taxation as their extractive profit is "
        "reallocated; once resource rights are granted prior to taxing rights they become net "
        "winners. Four formulas; gains at CIT, losses at ETR." + _CMN),
    6: ("Change in tax revenue across the three resource treatments (sales & employees; gains at "
        "CIT, losses at ETR)",
        "As Figure 4, for tax revenue. The '+ minimum royalty' bars add the IGF-ATAF "
        "minimum-royalty floor (hatched, in both the USD and the percentage panel) as a "
        "separate revenue stream on top of the unitary-taxation yield." + _REVPCT + _CMN),
    7: ("Guinea and Mali: the minimum royalty is decisive for under-capturing producers",
        "Guinea — a byword for bauxite rent under-capture — is a net revenue loser under all "
        "four formulas until the minimum royalty is enforced, which turns it into a clear "
        "winner. Mali, similarly emblematic for gold, gains under unitary taxation already, "
        "but the royalty lifts its gains by roughly 40 per cent under every formula. Gains at "
        "CIT, losses at ETR." + _CMN),
    8: ("Change in taxable profits by income group: origin vs destination-based sales",
        "Destination-based sales (all multinational destination-based sales plus the "
        "multinational share of digitally deliverable services imports), its nexus-adjusted "
        "variant, and origin-based sales as reported; preferred sales & employees formula, "
        "resource rights prior to taxing rights." + _POSNOTE + _CMN),
    9: ("Change in tax revenue by income group: origin vs destination-based sales",
        "The same three sales-factor measures as the previous figure, for tax revenue (gains "
        "at CIT, losses at ETR). Low- and lower-middle-income countries gain most under the "
        "destination-based measures." + _REVPCT + _CMN),
    10: ("Who gains from each apportionment factor? Change in taxable profit if profits were "
         "apportioned by one factor alone",
         "Estimated change in taxable profits per income group if multinational profits were "
         "apportioned by a single factor (left: US$ bn per year; right: % of the group's current "
         "profit base). Because the formulary allocation is linear in the factor weights, any "
         "formula's outcome is approximately the weight-average of these columns. Consistent "
         "with Loretz (2026), heavier weight on employment shifts taxing rights toward "
         "labour-intensive economies, destination-based sales benefit every non-haven group, "
         "and investment hubs lose under every factor. Resource rights prior to taxing rights."
         + _CMN),
    11: ("Revenue effect of cross-border loss consolidation, by income group",
         "The consolidation-adjusted bars credit zero tax on reported losses; because the "
         "consolidation loss is an upper bound, the adjusted revenue is a lower bound, and much "
         "of the gap is transitory (losses would be carried forward). Additional estimate, not "
         "the headline. Preferred sales & employees formula; resources excluded; gains at CIT, "
         "losses at ETR." + _CMN),
}


def _is_heading(p):
    s = (p.style.name if p.style else "") or ""
    return s.startswith("Heading") or s == "Title"


def find_heading(doc, *needles):
    for require_heading in (True, False):
        for p in doc.paragraphs:
            t = p.text.strip()
            if require_heading and not _is_heading(p):
                continue
            if any(t.startswith(n) or n in t for n in needles):
                return p
    return None


def find_caption(doc, fragment):
    """Caption anchor: the first paragraph that STARTS with 'Figure' and CONTAINS
    the fragment — robust to the author renumbering figures. Main-text captions
    precede same-titled Appendix F headings in document order, so first-match is
    the main-text one."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Figure") and fragment in t:
            return p
    return None


def _fmt(v):
    if isinstance(v, float):
        return f"{v:,.2f}"
    return "" if (v is None or (isinstance(v, float) and pd.isna(v))) else str(v)


def add_par_after(doc, anchor_el, text, bold=False):
    p = doc.add_paragraph()
    p.add_run(text).bold = bold
    anchor_el.addnext(p._p)
    return p._p


def add_image_after(doc, anchor_el, img_path, width_in=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    anchor_el.addnext(p._p)
    return p._p


def _remove_images_before(anchor_p):
    """Remove image paragraph(s) sitting directly above `anchor_p` (older pasted figures)."""
    prev = anchor_p._p.getprevious()
    while (prev is not None and prev.tag == qn("w:p")
           and prev.findall(".//" + qn("a:blip"))):
        rem, prev = prev, prev.getprevious()
        rem.getparent().remove(rem)


def _strip_drawings(p):
    """Remove inline images embedded in the paragraph itself (author-pasted figures
    often share the paragraph with the caption text)."""
    for tag in ("w:drawing", "w:pict"):
        for el in p._p.findall(".//" + qn(tag)):
            # remove the whole run holding the drawing if possible, else the drawing
            parent = el.getparent()
            while parent is not None and parent.tag != qn("w:r"):
                parent = parent.getparent()
            (parent if parent is not None else el).getparent().remove(
                parent if parent is not None else el)


def replace_images_under(doc, heading, img_path):
    """Remove image paragraph(s) between `heading` and the next heading, then insert
    the current image right after the heading. Non-image prose is left untouched."""
    el = heading._p.getnext()
    rem = []
    while el is not None:
        if el.tag == qn("w:p") and _is_heading_el(el):
            break
        if el.tag == qn("w:p") and (el.findall(".//" + qn("w:drawing"))
                                    or el.findall(".//" + qn("w:pict"))):
            rem.append(el)
        el = el.getnext()
    for e in rem:
        e.getparent().remove(e)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(6.5))
    heading._p.addnext(p._p)
    return p


def add_par_before(doc, anchor_p, text, bold=False, size=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    anchor_p._p.addprevious(p._p)
    return p


def add_image_before(doc, anchor_p, img_path, width_in=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    anchor_p._p.addprevious(p._p)
    return p


def add_fig_block_after(doc, anchor_el, title, img_path, note, num_label, width_in=6.5):
    """Append a title / image / note block right after `anchor_el` (used in appendices
    where there is no caption anchor). Returns the last inserted element."""
    t = doc.add_paragraph(); tr = t.add_run(f"{num_label} {title}"); tr.bold = True
    anchor_el.addnext(t._p); cur = t._p
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(img_path, width=Inches(width_in))
    cur.addnext(ip._p); cur = ip._p
    n = doc.add_paragraph(); nr = n.add_run(note); nr.italic = True; nr.font.size = Pt(8)
    cur.addnext(n._p)
    return n._p


def _shade(cell, fill):
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell(cell, text, bold=False, right=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_table_from_csv_after(doc, anchor_el, csv_path):
    """Insert a CSV as a Word table: bold shaded header, bold shaded group-heading
    rows (uppercase first cell), right-aligned numeric cells."""
    df = pd.read_csv(csv_path)
    headers = [str(c) for c in df.columns]
    tbl = doc.add_table(rows=1, cols=len(headers))
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for j, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[j], h, bold=True, right=(j > 0))
        _shade(tbl.rows[0].cells[j], "D9D9D9")
    for _, r in df.iterrows():
        cells = tbl.add_row().cells
        first = str(r.iloc[0])
        is_group = first.strip() != "" and first.strip() == first.strip().upper()
        for j, v in enumerate(r):
            _set_cell(cells[j], _fmt(v), bold=is_group, right=(j > 0))
            if is_group:
                _shade(cells[j], "F2F2F2")
    anchor_el.addnext(tbl._tbl)
    return tbl._tbl, len(df)


def insert_text(doc, heading, paragraphs):
    cur = add_par_after(doc, heading._p, MARK, bold=True)
    for t in paragraphs:
        cur = add_par_after(doc, cur, t)
    return cur


def _cap_caption(s):
    """Proper table heading: capitalise the leading 'table' and the first description word."""
    import re
    s = s[:1].upper() + s[1:] if s else s
    return re.sub(r"(\.\s+)([a-z])", lambda m: m.group(1) + m.group(2).upper(), s, count=1)


def insert_tables(doc, heading, items):
    cur = add_par_after(doc, heading._p, MARK, bold=True)
    for caption, fname in items:
        cur = add_par_after(doc, cur, _cap_caption(caption), bold=True)
        el, n = add_table_from_csv_after(doc, cur, os.path.join(TBL, fname))
        print(f"    + table {fname} ({n} rows)")
        cur = el
    return cur


def _is_heading_el(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return False
    ps = pPr.find(qn("w:pStyle"))
    return ps is not None and "Heading" in (ps.get(qn("w:val")) or "")


def replace_table_under(doc, heading, items):
    """Remove any existing table(s) between `heading` and the next heading, then insert the
    current table(s) right after the heading. items = [(sub_label_or_None, csv_filename), …].
    The heading itself is the caption, so no extra caption is added (matches the existing
    Appendix→Tables design: grey bold header, right-aligned numbers, shaded group rows)."""
    el, rem = heading._p.getnext(), []
    while el is not None:
        if el.tag == qn("w:tbl"):
            rem.append(el)
        elif el.tag == qn("w:p") and _is_heading_el(el):
            break
        el = el.getnext()
    for t in rem:
        t.getparent().remove(t)
    cur = heading._p
    for label, fname in items:
        if label:
            cur = add_par_after(doc, cur, label, bold=True)
        el2, n = add_table_from_csv_after(doc, cur, os.path.join(TBL, fname))
        print(f"    + {fname} ({n} rows) → '{heading.text.strip()[:34]}'")
        cur = el2
    return cur


def main():
    if not os.path.exists(ORIG):
        print(f"[abort] main doc not found: {ORIG}")
        return
    orig_mtime = os.path.getmtime(ORIG)
    target = None
    for n in range(4, 40):
        cand = os.path.join(_DIR, f"Unitary taxation_autodraft_v{n}.docx")
        try:
            shutil.copy2(ORIG, cand)
            target = cand
            break
        except PermissionError:
            print(f"  [skip] {os.path.basename(cand)} is open/locked")
    if target is None:
        print("[abort] no free autodraft target")
        return
    doc = docx.Document(target)

    # ── Main-text figures: refresh the IMAGE at each caption, keep author text ──
    n_fig = 0
    for needle, fname in FIG_CAPTIONS:
        img = os.path.join(FIG_REP, fname)
        cap = find_caption(doc, needle)
        if cap is None or not os.path.exists(img):
            print(f"  [warn] caption anchor or image missing: '{needle[:50]}' / {fname}")
            continue
        _remove_images_before(cap)      # bare image paragraph(s) above the caption
        _strip_drawings(cap)            # image pasted into the caption paragraph
        add_image_before(doc, cap, img)
        n_fig += 1
    print(f"  refreshed {n_fig}/{len(FIG_CAPTIONS)} main-text figure images "
          f"(captions untouched)")

    # ── "Tables" section: replace the outdated tables under each heading ─────
    # (Headings renumbered by the author 2026-07-12: Table 1 = taxbase by formula,
    # Table 2 = revenue by formula, Table 3 = break-even, Table 4 = scenarios.)
    for needles, items in [
        (("Table 1: Change in taxable profits",),
         [(None, "table2a_taxbase_by_formula__reported_only.csv")]),
        (("Table 2: Revenue estimates",),
         [(None, "table2b_revenue_by_formula__reported_only.csv")]),
        (("Table 3: Break-even",),
         [("Tax havens (average ETR):",
           "table5_breakeven_etr__reported_only.csv"),
          ("Home-bias countries (domestic ETR):",
           "table6_breakeven_domestic_etr__reported_only.csv")]),
        (("Table 4: Unitary taxation effects",),
         [(None, "table1_scenarios__reported_only.csv")]),
    ]:
        h = find_heading(doc, *needles)
        if h:
            replace_table_under(doc, h, items)
        else:
            print(f"  [warn] Tables heading not found: {needles[0]}")

    # ── Appendix D — dividend double-counting prose (only if not merged yet) ──
    if find_heading(doc, "Tax-haven list") is None:
        h = find_heading(doc, "Appendix D: Correcting country by country data")
        if h:
            insert_text(doc, h, TXT_APP_C)
            print("  + Appendix D dividend prose")
    else:
        print("  = Appendix D dividend prose already merged — skipped")

    # ── Appendix F — gravity results: intro prose (guarded), tables, figures ──
    h = find_heading(doc, "Appendix F: Results with disaggregated rows")
    if h:
        if not any("reproduces the main-text results" in p.text for p in doc.paragraphs):
            cur = add_par_after(doc, h._p, MARK, bold=True)
            for t in TXT_APP_E:
                cur = add_par_after(doc, cur, t)
            print("  + Appendix F intro prose")
        for needles, items in [
            (("Table F1.",), [(None, "table1_scenarios__gravity.csv")]),
            (("Table F2.",), [(None, "table2a_taxbase_by_formula__gravity.csv")]),
            (("Table F3.",), [(None, "table2b_revenue_by_formula__gravity.csv")]),
        ]:
            th = find_heading(doc, *needles)
            if th:
                replace_table_under(doc, th, items)
            else:
                print(f"  [warn] Appendix F heading not found: {needles[0]}")
        for needle, fname, retitle in FIG_APPENDIX_F:
            img = os.path.join(FIG_GRAV, fname)
            fh = find_heading(doc, needle)
            if fh is None or not os.path.exists(img):
                print(f"  [warn] Appendix F figure slot/image missing: {needle} / {fname}")
                continue
            if retitle:
                for r in list(fh.runs):
                    r.text = ""
                (fh.runs[0] if fh.runs else fh.add_run()).text = retitle
            replace_images_under(doc, fh, img)
        print("  + Appendix F gravity tables + figures refreshed")
    else:
        print("  [warn] 'Appendix F: Results with disaggregated rows' heading not found")

    # ── Appendix G — hand-collected sources (prose guarded; table replaced) ──
    h = find_heading(doc, "Appendix G: Overview of hand-collected")
    if h:
        # table first (anchored at the heading), then prose inserted directly
        # after the heading so it ends up ABOVE the table
        replace_table_under(doc, h, [
            (_cap_caption(CAP["fF"]), "appendixF_handcollected_sources.csv")])
        print("  + Appendix G hand-collected table")
        if not any("collected by hand" in p.text for p in doc.paragraphs):
            cur = add_par_after(doc, h._p, MARK, bold=True)
            for t in TXT_APP_F:
                cur = add_par_after(doc, cur, t)
            print("  + Appendix G prose")

    # ── Expanded references (flagged block after the References heading) ─────
    h = find_heading(doc, "References")
    if h:
        cur = add_par_after(doc, h._p, MARK + "  (merge into the list & verify)", bold=True)
        for t in REFS_ADD:
            cur = add_par_after(doc, cur, t, bold=t.startswith("—"))
        print(f"  + {len(REFS_ADD)} reference lines (flagged)")

    doc.save(target)
    print(f"\nSaved: {target}")
    # The builder only reads ORIG and writes the copy. If ORIG's mtime changed it is
    # because the author saved it externally while this ran (e.g. copying content in) —
    # our script never writes to it. Warn rather than crash.
    if abs(os.path.getmtime(ORIG) - orig_mtime) >= 1e-6:
        print(f"[note] main doc mtime changed during the run (external save) — the builder "
              f"did not write to it; only {os.path.basename(target)} was written.")
    else:
        print(f"Main doc unchanged: {ORIG}")


if __name__ == "__main__":
    main()
